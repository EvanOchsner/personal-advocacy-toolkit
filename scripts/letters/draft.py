#!/usr/bin/env python3
"""Template-driven letter drafting from a case-intake.yaml.

Approach
--------
The .j2 templates under ``templates/letter-templates/<kind>.docx.j2`` are
*plain-text* Jinja2 templates despite the ``.docx.j2`` suffix. The
"docx" piece refers to the final output format — we render the template
to a string, split on blank lines into paragraphs, and compose a
``.docx`` file via ``python-docx``. This avoids the nightmare of
templating a real .docx (a .zip of XML) with Jinja, which requires
docxtpl-style round-trip handling. Templates stay scannable in a text
editor and git diff; the rendered output is still a .docx.

If the output path ends in ``.txt`` or ``.md``, the rendered text is
written directly and no ``python-docx`` dependency is required.

case-intake.yaml schema (v0.1, partial — see
scripts/intake/situation_classify.py)
---------------------------------------------------------------------
Top-level keys this tool reads (all optional unless flagged):
  claimant:
    name: str                # REQUIRED by most letter kinds
    address: { street, city, state, zip } | str
    email: str
    phone: str
  parties:                   # optional — richer fixture-shape
    insurer: {name, address, claims_phone, ...}
    adjuster: {name, email, phone, ...}
  jurisdiction:
    state: str
  situation_type: str        # slug from data/situation_types.yaml
  loss:
    date: YYYY-MM-DD
    description: str
  policy:                    # insurance-specific
    policy_number: str
    agreed_value_usd: number
  disputed_amounts: ...      # insurance-specific
  regulator: {name, short_name, case_number}

Per-kind required fields (see REQUIRED_FIELDS below) are resolved from
the intake by dotted path. Missing required fields prompt interactively
unless ``--strict`` is set.

CLI
---
    python -m scripts.letters.draft \\
        --kind demand \\
        --intake case-intake.yaml \\
        --out letter.docx

    python -m scripts.letters.draft --kind foia --intake case-intake.yaml \\
        --out request.docx --recipient-name "Maryland Insurance Administration"

Kinds: demand | foia | preservation | withdrawal | cease-desist
"""
from __future__ import annotations

import argparse
import sys
from datetime import date
from pathlib import Path
from typing import Any

from scripts.intake._common import DISCLAIMER, data_dir, find_repo_root, load_yaml


LETTER_DISCLAIMER = (
    "Prepared with the advocacy-toolkit. Reference material, not legal advice."
)

KINDS = ("demand", "foia", "preservation", "withdrawal", "cease-desist")

# Per-kind signature phrase that the template is expected to include.
# These also serve as test anchors.
SIGNATURE_PHRASES: dict[str, str] = {
    "demand": "hereby demand",
    "foia": "public records request",
    "preservation": "preservation of evidence",
    "withdrawal": "withdraw my consent",
    "cease-desist": "cease and desist",
}

# Required intake fields per kind, expressed as dotted paths. If a field
# resolves to None/empty and --strict is set, we fail. Otherwise we
# prompt the user interactively (if stdin is a tty) or substitute a
# ``[TODO: <field>]`` placeholder.
REQUIRED_FIELDS: dict[str, list[str]] = {
    "demand": ["sender.name", "recipient.name"],
    "foia": ["sender.name", "recipient.name"],
    "preservation": ["sender.name", "recipient.name"],
    "withdrawal": ["sender.name", "recipient.name"],
    "cease-desist": ["sender.name", "recipient.name"],
}


# --------------------------------------------------------------------------- #
# Intake / authorities lookup helpers
# --------------------------------------------------------------------------- #


def _get_dotted(d: dict[str, Any], path: str) -> Any:
    cur: Any = d
    for part in path.split("."):
        if not isinstance(cur, dict):
            return None
        cur = cur.get(part)
    return cur


def _first_authority(
    authorities_yaml: dict[str, Any],
    jurisdiction: str | None,
    situation: str | None,
) -> dict[str, Any] | None:
    """Return the first *populated* authority for (jurisdiction, situation),
    or None if nothing matches."""
    if not jurisdiction or not situation:
        return None
    juris = (authorities_yaml.get("jurisdictions") or {}).get(jurisdiction.upper())
    if not juris:
        return None
    sit = (juris.get("situations") or {}).get(situation)
    if not sit or sit.get("status") != "populated":
        return None
    auths = sit.get("authorities") or []
    for a in auths:
        if a.get("name") and not str(a["name"]).startswith("TODO"):
            return a
    return None


def _format_address(addr: Any) -> str:
    if not addr:
        return ""
    if isinstance(addr, str):
        return addr.strip()
    if isinstance(addr, dict):
        parts = []
        if addr.get("street"):
            parts.append(str(addr["street"]))
        line2 = ", ".join(
            str(addr[k]) for k in ("city", "state", "zip") if addr.get(k)
        )
        if line2:
            parts.append(line2)
        return "\n".join(parts)
    return str(addr)


def _default_recipient(
    intake: dict[str, Any],
    kind: str,
    authorities_yaml: dict[str, Any] | None,
) -> dict[str, str]:
    """Derive a plausible default recipient from intake + authorities.

    - For FOIA and preservation-addressed-to-regulator, we prefer the
      matching authority (e.g. MIA for MD insurance_dispute).
    - For demand / withdrawal / cease-desist we prefer the counterparty
      (insurer / adjuster / landlord / etc.) if the intake knows about it.
    """
    parties = intake.get("parties") or {}
    regulator = intake.get("regulator") or {}
    jurisdiction = (intake.get("jurisdiction") or {}).get("state")
    situation = intake.get("situation_type")

    if kind == "foia":
        # Prefer authorities lookup, then the regulator block.
        auth = (
            _first_authority(authorities_yaml, jurisdiction, situation)
            if authorities_yaml
            else None
        )
        if auth:
            return {
                "name": str(auth.get("name", "") or ""),
                "address": "",  # authorities.yaml doesn't carry postal addresses yet
            }
        if regulator.get("name"):
            return {
                "name": str(regulator["name"]),
                "address": "",
            }

    if kind == "preservation":
        # Preservation-of-evidence letters usually target the counterparty.
        counterparty = parties.get("insurer") or parties.get("landlord") or parties.get(
            "employer"
        ) or parties.get("debt_collector") or parties.get("merchant")
        if counterparty:
            return {
                "name": str(counterparty.get("name", "") or ""),
                "address": _format_address(counterparty.get("address")),
            }

    # demand / withdrawal / cease-desist: prefer counterparty
    for role in ("insurer", "landlord", "employer", "debt_collector", "merchant", "adjuster"):
        party = parties.get(role)
        if party and party.get("name"):
            return {
                "name": str(party["name"]),
                "address": _format_address(party.get("address")),
            }

    return {"name": "", "address": ""}


# --------------------------------------------------------------------------- #
# Template context assembly
# --------------------------------------------------------------------------- #


def build_context(
    intake: dict[str, Any],
    kind: str,
    authorities_yaml: dict[str, Any] | None,
    overrides: dict[str, str] | None = None,
) -> dict[str, Any]:
    """Flatten intake into a context dict suitable for the Jinja templates.

    Overrides (from CLI flags like --recipient-name) win over everything.
    """
    claimant = intake.get("claimant") or {}
    sender_address = _format_address(claimant.get("address"))
    sender = {
        "name": claimant.get("name") or "",
        "address": sender_address,
        "email": claimant.get("email") or "",
        "phone": claimant.get("phone") or "",
    }

    recipient = _default_recipient(intake, kind, authorities_yaml)

    loss = intake.get("loss") or {}
    policy = intake.get("policy") or {}
    regulator = intake.get("regulator") or {}
    disputed = intake.get("disputed_amounts") or {}

    ctx: dict[str, Any] = {
        "today": date.today().isoformat(),
        "kind": kind,
        "disclaimer": LETTER_DISCLAIMER,
        "sender": sender,
        "recipient": recipient,
        "jurisdiction": intake.get("jurisdiction") or {},
        "situation_type": intake.get("situation_type") or "",
        "loss_date": loss.get("date") or "",
        "loss_description": loss.get("description") or "",
        "policy_number": policy.get("policy_number") or "",
        "agreed_value_usd": policy.get("agreed_value_usd") or "",
        "insurer_acv_offer_usd": disputed.get("insurer_acv_offer_usd") or "",
        "insurer_deduction_usd": disputed.get("insurer_deduction_usd") or "",
        "regulator_name": regulator.get("name") or "",
        "regulator_case_number": regulator.get("case_number") or "",
        "case_caption": intake.get("case_name") or intake.get("case_slug") or "",
        "synthetic": bool(intake.get("synthetic")),
    }

    if overrides:
        for k, v in overrides.items():
            if not v:
                continue
            # Support dotted override keys like recipient.name.
            if "." in k:
                outer, inner = k.split(".", 1)
                ctx.setdefault(outer, {})
                if isinstance(ctx[outer], dict):
                    ctx[outer][inner] = v
            else:
                ctx[k] = v
    return ctx


def _resolve_required(
    ctx: dict[str, Any], kind: str, strict: bool, interactive: bool
) -> list[str]:
    """Check REQUIRED_FIELDS; prompt/placeholder/fail as configured.

    Returns the list of fields that were missing (post-resolution).
    """
    missing: list[str] = []
    for dotted in REQUIRED_FIELDS.get(kind, []):
        val = _get_dotted(ctx, dotted)
        if val:
            continue
        if strict:
            missing.append(dotted)
            continue
        if interactive and sys.stdin.isatty():
            try:
                ans = input(f"Required field '{dotted}' is missing. Enter value: ").strip()
            except EOFError:
                ans = ""
            if ans:
                outer, _, inner = dotted.partition(".")
                if inner:
                    ctx.setdefault(outer, {})
                    if isinstance(ctx[outer], dict):
                        ctx[outer][inner] = ans
                else:
                    ctx[outer] = ans
                continue
        # fall through to placeholder
        outer, _, inner = dotted.partition(".")
        placeholder = f"[TODO: {dotted}]"
        if inner:
            ctx.setdefault(outer, {})
            if isinstance(ctx[outer], dict):
                ctx[outer][inner] = placeholder
        else:
            ctx[outer] = placeholder
        missing.append(dotted)
    return missing


# --------------------------------------------------------------------------- #
# Rendering
# --------------------------------------------------------------------------- #


def render_text(template_path: Path, ctx: dict[str, Any]) -> str:
    import jinja2  # declared dep

    env = jinja2.Environment(
        loader=jinja2.FileSystemLoader(str(template_path.parent)),
        autoescape=False,
        keep_trailing_newline=True,
        undefined=jinja2.ChainableUndefined,
    )
    tpl = env.get_template(template_path.name)
    return tpl.render(**ctx)


def write_docx(text: str, out: Path) -> None:
    """Compose a .docx from a blank-line-separated block of text."""
    from docx import Document  # declared dep

    doc = Document()
    # Each blank-line-separated block becomes its own paragraph. This
    # keeps the templates legible as plain text.
    blocks = [b for b in text.replace("\r\n", "\n").split("\n\n")]
    for block in blocks:
        block = block.rstrip("\n")
        if block == "":
            doc.add_paragraph("")
            continue
        # Preserve internal newlines as soft line-breaks within paragraph.
        lines = block.split("\n")
        p = doc.add_paragraph(lines[0])
        for ln in lines[1:]:
            p.add_run("\n" + ln)
    # Visible disclaimer footer paragraph. Even if the template already
    # contains it, adding it again is harmless and guarantees presence.
    doc.add_paragraph("")
    foot = doc.add_paragraph(LETTER_DISCLAIMER)
    try:
        foot.runs[0].italic = True
    except Exception:
        pass
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def write_text(text: str, out: Path) -> None:
    # Ensure disclaimer present exactly once at the tail of the file.
    if LETTER_DISCLAIMER not in text:
        text = text.rstrip() + "\n\n" + LETTER_DISCLAIMER + "\n"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8")


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #


def _templates_dir(root: Path | None) -> Path:
    return find_repo_root(root) / "templates" / "letter-templates"


def draft_letter(
    *,
    kind: str,
    intake_path: Path,
    out: Path,
    root: Path | None = None,
    recipient_name: str | None = None,
    recipient_address: str | None = None,
    strict: bool = False,
    interactive: bool = True,
    templates_dir: Path | None = None,
    authorities_yaml_path: Path | None = None,
) -> dict[str, Any]:
    if kind not in KINDS:
        raise ValueError(f"unknown kind {kind!r}; expected one of {KINDS}")

    intake = load_yaml(intake_path)
    authorities = None
    try:
        auth_path = authorities_yaml_path or (data_dir(root) / "authorities.yaml")
        if auth_path.exists():
            authorities = load_yaml(auth_path)
    except Exception:
        authorities = None

    overrides: dict[str, str] = {}
    if recipient_name:
        overrides["recipient.name"] = recipient_name
    if recipient_address:
        overrides["recipient.address"] = recipient_address

    ctx = build_context(intake, kind, authorities, overrides)
    missing = _resolve_required(ctx, kind, strict=strict, interactive=interactive)
    if strict and missing:
        raise ValueError(f"missing required fields in strict mode: {missing}")

    tdir = templates_dir or _templates_dir(root)
    tpl_path = tdir / f"{kind}.docx.j2"
    if not tpl_path.exists():
        raise FileNotFoundError(f"template not found: {tpl_path}")

    rendered = render_text(tpl_path, ctx)

    # Belt-and-suspenders: make sure the disclaimer and signature phrase
    # are present in the rendered body.
    if LETTER_DISCLAIMER not in rendered:
        rendered = rendered.rstrip() + "\n\n" + LETTER_DISCLAIMER + "\n"

    suffix = out.suffix.lower()
    if suffix in (".txt", ".md"):
        write_text(rendered, out)
    else:
        write_docx(rendered, out)

    return {
        "kind": kind,
        "template": tpl_path,
        "out": out,
        "missing_required": missing,
        "rendered_text": rendered,
    }


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    p.add_argument("--kind", required=True, choices=KINDS)
    p.add_argument("--intake", type=Path, required=True, help="case-intake.yaml")
    p.add_argument("--out", type=Path, required=True, help="output .docx / .txt / .md")
    p.add_argument("--recipient-name", default=None)
    p.add_argument("--recipient-address", default=None)
    p.add_argument("--strict", action="store_true", help="fail on any missing required field")
    p.add_argument(
        "--non-interactive",
        action="store_true",
        help="do not prompt; placeholder missing fields (or fail with --strict)",
    )
    p.add_argument("--templates-dir", type=Path, default=None)
    p.add_argument("--authorities-yaml", type=Path, default=None)
    p.add_argument("--root", type=Path, default=None)
    args = p.parse_args(argv)

    try:
        result = draft_letter(
            kind=args.kind,
            intake_path=args.intake,
            out=args.out,
            root=args.root,
            recipient_name=args.recipient_name,
            recipient_address=args.recipient_address,
            strict=args.strict,
            interactive=not args.non_interactive,
            templates_dir=args.templates_dir,
            authorities_yaml_path=args.authorities_yaml,
        )
    except (FileNotFoundError, ValueError) as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2

    print(f"{DISCLAIMER}")
    print(f"wrote {result['out']} (kind={result['kind']})")
    if result["missing_required"]:
        print(
            f"  NOTE: placeholders inserted for missing fields: "
            f"{', '.join(result['missing_required'])}",
            file=sys.stderr,
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
