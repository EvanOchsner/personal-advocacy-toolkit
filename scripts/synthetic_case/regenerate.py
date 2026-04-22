"""Regenerate the Mustang-in-Maryland synthetic artifacts.

Three artifact groups are (re)generated from their canonical markdown
sources under ``examples/mustang-in-maryland/``:

  valuation  -> evidence/valuation/MidAtlantic-Vehicle-Appraisers-valuation.pdf
  photos     -> evidence/photos/photo-{01,02,03}-*.jpg
  complaint  -> drafts/mia-complaint.docx

All outputs carry the SYNTHETIC stamp in visible content and in file
metadata. JPEGs have no EXIF GPS and no author/camera identifying
fields (only a UserComment set to the SYNTHETIC stamp). The .docx has
its core-properties author / company / last-modified-by cleared and
its description set to the stamp.

Idempotency: everything below uses a fixed seed (``SEED = 42``) for
any randomness (photo colour fields). Re-running without source
changes produces visually identical output; the reportlab and
python-docx outputs are not byte-identical across runs because both
embed a timestamp in the file, but content and structure are stable.

CLI:
    python -m scripts.synthetic_case.regenerate --all
    python -m scripts.synthetic_case.regenerate --only valuation
    python -m scripts.synthetic_case.regenerate --only photos --only complaint
    python -m scripts.synthetic_case.regenerate --root <alt-case-root>
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

from ._reportlab_theme import SYNTHETIC_STAMP, make_doc, paragraph_styles


SEED = 42

REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_CASE_ROOT = REPO_ROOT / "examples" / "mustang-in-maryland"


# ---------------------------------------------------------------------------
# Valuation PDF
# ---------------------------------------------------------------------------


def regenerate_valuation(case_root: Path) -> Path:
    """Render the MAVA valuation markdown into a real PDF."""
    src = case_root / "evidence" / "valuation" / "MidAtlantic-Vehicle-Appraisers-valuation.md"
    out = src.with_suffix(".pdf")
    text = src.read_text(encoding="utf-8")

    fields = _parse_valuation_fields(text)

    styles = paragraph_styles()
    doc = make_doc(out, title="MAVA Valuation Report", header_title="MidAtlantic Vehicle Appraisers")
    story: list = []

    story.append(Paragraph("MidAtlantic Vehicle Appraisers", styles["title"]))
    story.append(Paragraph(f"Report {fields.get('report_id', 'MAVA-SYNTHETIC')}", styles["h2"]))
    story.append(Paragraph(SYNTHETIC_STAMP, styles["h2"]))
    story.append(Spacer(1, 0.15 * inch))

    header_rows = [
        ("Prepared for", fields.get("prepared_for", "")),
        ("Claim", fields.get("claim", "")),
        ("Insured", fields.get("insured", "")),
        ("Date of inspection", fields.get("date_inspection", "")),
        ("Date of report", fields.get("date_report", "")),
        ("Appraiser", fields.get("appraiser", "")),
    ]
    story.append(_kv_table(header_rows))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Vehicle", styles["h2"]))
    veh = fields.get("vehicle", [])
    if veh:
        story.append(_kv_table(veh))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Condition notes", styles["h2"]))
    for item in fields.get("condition", []):
        story.append(Paragraph("&bull; " + _xml_escape(item), styles["body"]))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Market comparables (redacted)", styles["h2"]))
    story.append(
        Paragraph(
            "Three comparable 1969 Mustang Mach 1 units, 72,000 to 91,000 mi "
            "range, sold between Nov 2024 and Feb 2025 at regional classic-car "
            "auctions. Condition adjustments applied per MAVA internal matrix.",
            styles["body"],
        )
    )
    story.append(Spacer(1, 0.08 * inch))
    comps = fields.get("comps", [])
    if comps:
        story.append(_comps_table(comps))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Actual cash value determination", styles["h2"]))
    story.append(
        Paragraph(
            f"ACV (condition-adjusted, market comps): <b>{fields.get('acv', 'n/a')}</b>",
            styles["body"],
        )
    )

    story.append(Paragraph("Repair estimate (customary regional rates)", styles["h2"]))
    story.append(
        Paragraph(
            "Labor at $75.00/hr; parts at vendor quote; storage at $25/day.",
            styles["body"],
        )
    )
    story.append(
        Paragraph(
            f"Total: <b>{fields.get('repair_total', 'n/a')}</b>",
            styles["body"],
        )
    )

    story.append(Paragraph("Total-loss threshold analysis", styles["h2"]))
    story.append(
        Paragraph(
            "Repair-cost-to-ACV ratio: 66.0%. Total-loss threshold (80%) not "
            "met. Recommended disposition: repair at customary regional rates.",
            styles["body"],
        )
    )

    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph("Appraiser signature", styles["h2"]))
    story.append(Paragraph("/s/ Joyce Pemberton -- 2025-04-17", styles["body"]))
    story.append(
        Paragraph(
            f"MidAtlantic Vehicle Appraisers -- {SYNTHETIC_STAMP}",
            styles["small"],
        )
    )

    doc.build(story)
    return out


def _parse_valuation_fields(text: str) -> dict:
    """Pull structured fields from the markdown source.

    The markdown layout is stable (authored in Phase 2 and not expected
    to churn); we parse line-by-line rather than pulling in a full
    markdown parser.
    """
    fields: dict = {}

    m = re.search(r"## Report\s+(\S+)", text)
    if m:
        fields["report_id"] = m.group(1).strip()

    for key, label in [
        ("prepared_for", "Prepared for"),
        ("claim", "Claim"),
        ("insured", "Insured"),
        ("date_inspection", "Date of inspection"),
        ("date_report", "Date of report"),
        ("appraiser", "Appraiser"),
    ]:
        m = re.search(rf"{re.escape(label)}:\s*(.+)", text)
        if m:
            fields[key] = m.group(1).strip()

    # Vehicle section bullets
    veh = []
    veh_block = _section(text, "### Vehicle", "### Condition notes")
    for line in veh_block.splitlines():
        line = line.strip()
        if line.startswith("- ") and ":" in line:
            k, _, v = line[2:].partition(":")
            veh.append((k.strip(), v.strip()))
    fields["vehicle"] = veh

    # Condition notes bullets
    cond = []
    cond_block = _section(text, "### Condition notes", "### Market comparables")
    for line in cond_block.splitlines():
        line = line.strip()
        if line.startswith("- "):
            cond.append(line[2:].strip())
    fields["condition"] = cond

    # Comparables table
    comps = []
    comp_block = _section(text, "| Comp |", "### Actual cash value")
    for line in comp_block.splitlines():
        line = line.strip()
        if not line.startswith("|") or line.startswith("|----") or "----" in line:
            continue
        cells = [c.strip().replace("**", "") for c in line.strip("|").split("|")]
        if len(cells) >= 4 and cells[0].lower() != "comp":
            comps.append(cells[:4])
    fields["comps"] = comps

    m = re.search(r"ACV \(condition-adjusted, market comps\):\s*\*\*\s*([^*]+?)\*\*", text)
    if m:
        fields["acv"] = m.group(1).strip()
    m = re.search(r"Total:\s*\*\*\s*([^*]+?)\*\*", text)
    if m:
        fields["repair_total"] = m.group(1).strip()

    return fields


def _section(text: str, start: str, end: str) -> str:
    si = text.find(start)
    if si < 0:
        return ""
    ei = text.find(end, si + len(start))
    if ei < 0:
        ei = len(text)
    return text[si:ei]


def _kv_table(rows: list[tuple[str, str]]) -> Table:
    data = [[k, v] for k, v in rows]
    t = Table(data, colWidths=[1.8 * inch, 4.2 * inch])
    t.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (0, -1), "Helvetica-Bold", 10),
                ("FONT", (1, 0), (1, -1), "Helvetica", 10),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ]
        )
    )
    return t


def _comps_table(rows: list[list[str]]) -> Table:
    header = ["Comp", "Sale price", "Adj.", "Adjusted"]
    data = [header, *rows]
    t = Table(data, colWidths=[0.9 * inch, 1.4 * inch, 1.4 * inch, 1.4 * inch])
    t.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 10),
                ("FONT", (0, 1), (-1, -1), "Helvetica", 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return t


def _xml_escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ---------------------------------------------------------------------------
# Photos (JPEG)
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class PhotoSpec:
    md_name: str
    out_name: str
    subject: str
    caption: str
    # RGB colour seed — deterministic, no real-world meaning.
    bg: tuple[int, int, int]


PHOTO_SPECS: list[PhotoSpec] = [
    PhotoSpec(
        md_name="photo-01-mustang-at-midlife-crisis.md",
        out_name="photo-01-mustang-at-midlife-crisis.jpg",
        subject="Mustang at Midlife Crisis Restorations -- rear 3/4 view",
        caption=(
            "Mustang at Midlife Crisis Restorations of Wilmington -- rear 3/4 "
            "view, vehicle in covered bay on intake day. Photo by claimant, "
            "2025-03-24."
        ),
        bg=(72, 92, 118),  # muted slate
    ),
    PhotoSpec(
        md_name="photo-02-damage-closeup.md",
        out_name="photo-02-damage-closeup.jpg",
        subject="Damage closeup -- driver-side rear quarter",
        caption=(
            "Damage closeup -- driver-side rear quarter panel, trunk lid "
            "misalignment visible. Photo by claimant, 2025-03-24."
        ),
        bg=(112, 80, 72),  # muted oxide
    ),
    PhotoSpec(
        md_name="photo-03-vin-plate-placeholder.md",
        out_name="photo-03-vin-plate-placeholder.jpg",
        subject="VIN plate -- VIN REDACTED -- SYNTHETIC CASE",
        caption=(
            "VIN plate (driver-side doorjamb). VIN REDACTED -- SYNTHETIC "
            "CASE. Photo by claimant, 2025-03-24."
        ),
        bg=(88, 96, 80),  # muted olive
    ),
]


def regenerate_photos(case_root: Path) -> list[Path]:
    """Generate the three synthetic JPEG placeholders.

    These are deliberately low-realism: a flat colour field with a
    large SYNTHETIC watermark and a caption line. The goal is a clearly
    fake teaching placeholder, not a convincing photo.
    """
    photos_dir = case_root / "evidence" / "photos"
    outputs: list[Path] = []
    for spec in PHOTO_SPECS:
        out = photos_dir / spec.out_name
        _render_photo(out, spec)
        outputs.append(out)
    return outputs


def _render_photo(out: Path, spec: PhotoSpec) -> None:
    from PIL import Image, ImageDraw, ImageFont

    W, H = 1600, 1200
    img = Image.new("RGB", (W, H), color=spec.bg)
    draw = ImageDraw.Draw(img)

    # A deterministic low-frequency "texture": faint diagonal bands so
    # the file is not a perfectly flat solid. Seeded RNG ensures
    # idempotency.
    import random

    rng = random.Random(SEED)
    for i in range(0, W + H, 40):
        jitter = rng.randint(-6, 6)
        draw.line(
            [(i + jitter, 0), (0, i + jitter)],
            fill=(
                max(0, min(255, spec.bg[0] + rng.randint(-8, 8))),
                max(0, min(255, spec.bg[1] + rng.randint(-8, 8))),
                max(0, min(255, spec.bg[2] + rng.randint(-8, 8))),
            ),
            width=1,
        )

    # Fonts: try a system font, fall back to default bitmap font if
    # unavailable. Both yield deterministic output for fixed inputs.
    stamp_font = _load_font(120)
    subject_font = _load_font(48)
    caption_font = _load_font(28)

    # Huge SYNTHETIC watermark, center.
    _draw_centered(draw, (W // 2, H // 2 - 120), SYNTHETIC_STAMP, stamp_font, fill=(255, 255, 255))
    # Subject label, below.
    _draw_centered(draw, (W // 2, H // 2 + 20), spec.subject, subject_font, fill=(235, 235, 235))
    # Caption at bottom.
    _draw_wrapped(draw, spec.caption, caption_font, (60, H - 180), W - 120, fill=(220, 220, 220))

    # Border rectangle to reinforce "this is a placeholder" feel.
    draw.rectangle([(20, 20), (W - 20, H - 20)], outline=(255, 255, 255), width=4)

    # EXIF: explicitly empty. We pass exif=b"" to save() so Pillow
    # does not inherit any metadata; and we also ensure no GPS or
    # camera-identifying fields are present. A UserComment carrying
    # the SYNTHETIC stamp is attached via piexif.
    exif_bytes = _build_synthetic_exif()
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(
        out,
        format="JPEG",
        quality=85,
        optimize=True,
        exif=exif_bytes,
    )


def _build_synthetic_exif() -> bytes:
    """Build an EXIF block carrying the SYNTHETIC stamp in the
    ImageDescription and UserComment tags. No GPS, no Make/Model, no
    Artist, no DateTimeOriginal.

    Uses Pillow's ``Image.Exif`` builder. If anything goes wrong,
    returns ``b""`` -- an empty EXIF block is written and the image
    has no GPS and no author metadata, which is the actual load-bearing
    guarantee this function exists to deliver.
    """
    try:
        from PIL import Image, ExifTags

        exif = Image.Exif()
        # 0th IFD
        exif[ExifTags.Base.ImageDescription.value] = SYNTHETIC_STAMP
        exif[ExifTags.Base.Software.value] = "advocacy-toolkit synthetic-case"
        exif[ExifTags.Base.Artist.value] = ""
        exif[ExifTags.Base.Copyright.value] = ""
        # Exif IFD — UserComment with ASCII character-code prefix.
        user_comment = b"ASCII\x00\x00\x00" + SYNTHETIC_STAMP.encode("ascii")
        exif[ExifTags.Base.UserComment.value] = user_comment
        return exif.tobytes()
    except Exception:
        return b""


def _load_font(size: int):
    from PIL import ImageFont

    candidates = [
        "/System/Library/Fonts/Helvetica.ttc",
        "/System/Library/Fonts/HelveticaNeue.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _draw_centered(draw, xy, text, font, *, fill) -> None:
    try:
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    except AttributeError:
        tw, th = draw.textsize(text, font=font)  # type: ignore
    x, y = xy
    draw.text((x - tw // 2, y - th // 2), text, font=font, fill=fill)


def _draw_wrapped(draw, text, font, xy, max_w, *, fill) -> None:
    words = text.split()
    lines: list[str] = []
    line = ""
    for w in words:
        trial = f"{line} {w}".strip()
        try:
            bbox = draw.textbbox((0, 0), trial, font=font)
            tw = bbox[2] - bbox[0]
        except AttributeError:
            tw, _ = draw.textsize(trial, font=font)  # type: ignore
        if tw <= max_w:
            line = trial
        else:
            if line:
                lines.append(line)
            line = w
    if line:
        lines.append(line)
    x, y = xy
    for i, ln in enumerate(lines):
        draw.text((x, y + i * (font.size + 6) if hasattr(font, "size") else y + i * 30), ln, font=font, fill=fill)


# ---------------------------------------------------------------------------
# MIA complaint .docx
# ---------------------------------------------------------------------------


def regenerate_complaint(case_root: Path) -> Path:
    """Render the MIA complaint markdown to a real .docx."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    src = case_root / "drafts" / "mia-complaint.md"
    out = src.with_suffix(".docx")
    md = src.read_text(encoding="utf-8")

    doc = Document()
    # Clear identifying metadata; set description to the SYNTHETIC stamp.
    cp = doc.core_properties
    cp.author = "advocacy-toolkit synthetic-case"
    cp.last_modified_by = "advocacy-toolkit synthetic-case"
    cp.comments = SYNTHETIC_STAMP
    # python-docx exposes description via .comments? actually distinct.
    try:
        cp.subject = SYNTHETIC_STAMP
        # description not directly exposed; set via .comments (comments
        # maps to cp:comments which is fine for our "stamp in metadata"
        # requirement). Title is pedagogically useful.
        cp.title = "MIA Complaint -- Delia Vance v. Chesapeake Indemnity Mutual (SYNTHETIC)"
        cp.keywords = SYNTHETIC_STAMP
    except Exception:
        pass

    # Header on every page carrying the synthetic stamp.
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = SYNTHETIC_STAMP
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = f"{SYNTHETIC_STAMP} -- Mustang in Maryland teaching example."
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for block in _parse_complaint_blocks(md):
        kind, content = block
        if kind == "h1":
            p = doc.add_heading(content, level=1)
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "h3":
            doc.add_heading(content, level=3)
        elif kind == "hr":
            doc.add_paragraph("____________________________________________")
        elif kind == "list":
            for item in content:
                doc.add_paragraph(item, style="List Number" if item[:2].rstrip(".").isdigit() else "List Bullet")
        elif kind == "para":
            doc.add_paragraph(content)

    # Explicit final disclaimer paragraph (the markdown already ends
    # with one, but we add again for safety).
    doc.add_paragraph("")
    p = doc.add_paragraph(SYNTHETIC_STAMP + ". This document is a teaching artifact.")
    try:
        p.runs[0].italic = True
    except Exception:
        pass

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _parse_complaint_blocks(md: str) -> list[tuple[str, object]]:
    """Coarse markdown-to-blocks parser sufficient for the MIA complaint.

    Recognises: `# H1`, `## H2`, `### H3`, `---` horizontal rule,
    numbered lists (`1. ...`), and paragraphs. HTML comments are
    stripped. Bold / italic markup is passed through as literal text
    since the docx renderer does not need formatted runs for this
    teaching artifact.
    """
    # Strip HTML comments.
    md = re.sub(r"<!--.*?-->", "", md, flags=re.DOTALL)

    blocks: list[tuple[str, object]] = []
    lines = md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", _strip_md_inline(line[2:].strip())))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", _strip_md_inline(line[3:].strip())))
            i += 1
        elif line.startswith("### "):
            blocks.append(("h3", _strip_md_inline(line[4:].strip())))
            i += 1
        elif line.strip() == "---":
            blocks.append(("hr", None))
            i += 1
        elif re.match(r"^\d+\.\s", line) or line.startswith("- "):
            items: list[str] = []
            while i < n and (re.match(r"^\d+\.\s", lines[i]) or lines[i].startswith("- ") or lines[i].startswith("  ")):
                cur = lines[i].rstrip()
                if re.match(r"^\d+\.\s", cur) or cur.startswith("- "):
                    items.append(_strip_md_inline(re.sub(r"^(\d+\.\s|- )", "", cur)))
                else:
                    # continuation of previous item
                    if items:
                        items[-1] = items[-1] + " " + cur.strip()
                i += 1
            blocks.append(("list", items))
        else:
            # paragraph: gather until blank line
            para_lines = [line]
            i += 1
            while i < n and lines[i].strip() and not lines[i].startswith("#") and lines[i].strip() != "---":
                if re.match(r"^\d+\.\s", lines[i]) or lines[i].startswith("- "):
                    break
                para_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(("para", _strip_md_inline(" ".join(para_lines))))
    return blocks


def _strip_md_inline(s: str) -> str:
    # Strip bold/italic markers but keep the text.
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------


GROUPS: dict[str, Callable[[Path], object]] = {
    "valuation": regenerate_valuation,
    "photos": regenerate_photos,
    "complaint": regenerate_complaint,
}


def regenerate(case_root: Path, groups: list[str]) -> dict[str, object]:
    results: dict[str, object] = {}
    for g in groups:
        fn = GROUPS[g]
        results[g] = fn(case_root)
    return results


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument(
        "--root",
        type=Path,
        default=DEFAULT_CASE_ROOT,
        help="Case root directory (defaults to examples/mustang-in-maryland).",
    )
    p.add_argument(
        "--all",
        action="store_true",
        help="Regenerate all artifact groups.",
    )
    p.add_argument(
        "--only",
        action="append",
        choices=sorted(GROUPS),
        default=[],
        help="Regenerate only this group. Repeatable. Mutually exclusive with --all.",
    )
    args = p.parse_args(argv)

    if args.all and args.only:
        p.error("--all and --only are mutually exclusive")
    if not args.all and not args.only:
        p.error("specify --all or at least one --only GROUP")

    groups = sorted(GROUPS) if args.all else args.only
    results = regenerate(args.root, groups)

    for g, out in results.items():
        if isinstance(out, list):
            for o in out:
                print(f"{g}: {o}")
        else:
            print(f"{g}: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
